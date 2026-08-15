from pathlib import Path
from datetime import datetime, date
import re


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_text(path: str) -> str:
    """
    Extract text from PDF, DOCX, TXT, or MD.
    """
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".pdf":
        import fitz

        doc = fitz.open(p)
        return "\n".join(page.get_text() for page in doc)

    if ext == ".docx":
        from docx import Document

        doc = Document(p)

        parts = [
            para.text
            for para in doc.paragraphs
            if para.text.strip()
        ]

        for table in doc.tables:
            for row in table.rows:
                parts.append(
                    " | ".join(cell.text for cell in row.cells)
                )

        return "\n".join(parts)

    if ext in {".txt", ".md"}:
        return p.read_text(
            encoding="utf-8",
            errors="ignore"
        )

    raise ValueError(f"Unsupported file type: {ext}")


# ============================================================
# TEXT NORMALIZATION
# ============================================================

def normalize_text(text: str) -> str:
    """
    Normalize text for matching while preserving the original
    text elsewhere.
    """
    text = text.lower()

    replacements = {
        "sketchup": "sketch up",
        "autocad": "auto cad",
        "3d visualization": "3d visualization",
        "3d modelling": "3d modeling",
        "modelling": "modeling",
        "interior designer": "interior design",
        "interior architect": "interior architecture",
        "architectural drafter": "drafter",
        "architectural drafting": "technical drawing",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text

SECTION_HEADERS = {
    # ========================================================
    # EDUCATION
    # ========================================================

    "education",
    "education level",
    "academic background",
    "academic qualification",
    "academic qualifications",
    "educational background",
    "educational qualification",
    "education qualification",

    "pendidikan",
    "riwayat pendidikan",
    "latar belakang pendidikan",

    
    # ========================================================
    # EXPERIENCE
    # ========================================================

    "experience",
    "experiences",
    "work experience",
    "work experiences",
    "professional experience",
    "professional experiences",
    "employment history",
    "career history",
    "work history",

    "pengalaman",
    "pengalaman kerja",
    "riwayat kerja",
    "riwayat karier",

    # ========================================================
    # ORGANIZATION
    # ========================================================

    "organizational experience",
    "organisational experience",
    "organizational experiences",
    "organisational experiences",
    "organization experience",
    "organisation experience",
    "organizational activities",
    "organisational activities",

    "pengalaman organisasi",
    "organisasi",
    "kegiatan organisasi",

    # ========================================================
    # SKILLS
    # ========================================================

    "skills",
    "skill",
    "technical skills",
    "technical skill",
    "soft skills",
    "soft skill",
    "hard skills",
    "hard skill",
    "soft and hard skills",
    "professional skills",
    "core skills",

    "keahlian",
    "keahlian teknis",
    "kemampuan",
    "keterampilan",

    # ========================================================
    # LANGUAGES
    # ========================================================

    "languages",
    "language",

    "bahasa",
    "bahasa yang dikuasai",

    # ========================================================
    # PROJECTS
    # ========================================================

    "projects",
    "project experience",
    "project experiences",
    "academic projects",
    "personal projects",

    "proyek",
    "pengalaman proyek",

    # ========================================================
    # CERTIFICATIONS / TRAINING
    # ========================================================

    "certifications",
    "certification",
    "certificates",
    "certificate",
    "certifications and training",
    "training",
    "courses",

    "sertifikasi",
    "sertifikat",
    "pelatihan",
    "kursus",

    # ========================================================
    # ACHIEVEMENTS
    # ========================================================

    "achievements",
    "achievement",
    "awards",
    "award",
    "honors",
    "honours",

    "pencapaian",
    "prestasi",
    "penghargaan",

    # ========================================================
    # ADDITIONAL INFORMATION
    # ========================================================

    "additional information",
    "additional details",
    "other information",
    "others",

    "informasi tambahan",

    # ========================================================
    # NON-FORMAL EDUCATION
    # ========================================================

    "non-formal education",
    "non formal education",
    "informal education",

    "pendidikan non formal",
    "pendidikan nonformal",

    # ========================================================
    # VOLUNTEER
    # ========================================================

    "volunteer experience",
    "volunteering",

    "pengalaman sukarelawan",
    "kegiatan sukarela",

    # ========================================================
    # PUBLICATIONS
    # ========================================================

    "publications",
    "publication",

    "publikasi",

    # ========================================================
    # PROFILE / SUMMARY
    # ========================================================

    "profile",
    "summary",
    "professional summary",
    "objective",
    "career objective",

    "profil",
    "ringkasan",
    "tujuan karier",

    
}

def _clean_header(line: str) -> str:
    """
    Normalize CV section headers.

    Examples:
        'WORK EXPERIENCE'
        'Work Experience'
        'W O R K E X P E R I E N C E'

    are normalized consistently.
    """

    line = line.lower().strip()

    # Remove decorative characters.
    line = re.sub(
        r"[^a-z0-9\s]",
        " ",
        line
    )

    # Collapse whitespace.
    line = re.sub(
        r"\s+",
        " ",
        line
    ).strip()

    return line

def _header_key(line: str) -> str:
    """
    Create a comparison key for CV section headers.

    This makes these equivalent:

        WORK EXPERIENCE
        W O R K E X P E R I E N C E

    Result:
        workexperience
    """

    cleaned = _clean_header(line)

    return re.sub(
        r"\s+",
        "",
        cleaned
    )

def _normalize_spaced_header(line: str) -> str:
    """
    Convert stylized PDF headings such as:

        W O R K E X P E R I E N C E

    into:

        work experience
    """

    cleaned = line.lower().strip()

    # If every character is separated by spaces,
    # collapse the individual letters.
    tokens = cleaned.split()

    if (
        len(tokens) >= 3
        and all(
            len(token) == 1
            for token in tokens
        )
    ):
        return "".join(tokens)

    return cleaned

def _section(text: str, names):
    """
    Extract a CV section using flexible header matching.

    Supports:
    - English / Indonesian headers
    - Different capitalization
    - Extra whitespace
    - PDF decorative spaced letters
    """

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    # Normalize target headers.
    target_headers = {
        _header_key(name)
        for name in names
    }

    # Normalize ALL known section headers.
    known_section_headers = {
        _header_key(header)
        for header in SECTION_HEADERS
    }

    start = None

    # --------------------------------------------------------
    # FIND START HEADER
    # --------------------------------------------------------

    for i, line in enumerate(lines):

        header_key = _header_key(line)

        if header_key in target_headers:
            start = i + 1
            break

    # Section not found.
    if start is None:
        return ""

    result = []

    # --------------------------------------------------------
    # COLLECT SECTION CONTENT
    # --------------------------------------------------------

    for line in lines[start:]:

        header_key = _header_key(line)

        # Stop when another known section starts.
        if header_key in known_section_headers:
            break

        result.append(line)

    return "\n".join(result).strip()


def extract_section(text: str, names) -> str:
    """
    Public wrapper around the internal section-extraction logic.

    Exists so other modules (e.g. matcher.py, which parses Job
    Descriptions using the same header-matching heuristics) don't
    need to reach into a private/underscored function.
    """
    return _section(text, names)


# ============================================================
# NAME EXTRACTION
# ============================================================

def extract_name(
    text: str,
    fallback: str = "Unknown Candidate"
):
    """
    Extract candidate name from the top/header area of a CV.

    Most CV formats place the candidate name at the beginning.
    Jobstreet-generated CVs may contain a platform label before
    the actual candidate name.
    """

    lines = [
        re.sub(r"\s+", " ", line.strip())
        for line in text.splitlines()
        if line.strip()
    ]

    if not lines:
        return fallback

    ignored_labels = {
        "dibuat dengan profil jobstreet",
        "made with jobstreet profile",
        "curriculum vitae",
        "resume",
        "cv",
    }

    section_headers = {
        _header_key(header)
        for header in SECTION_HEADERS
    }

    # --------------------------------------------------------
    # Only inspect the CV header.
    # --------------------------------------------------------

    header_lines = lines[:10]

    for line in header_lines:

        clean = line.strip()
        low = clean.lower()

        # Skip known platform/document labels.
        if low in ignored_labels:
            continue

        # Skip section headers.
        if _header_key(clean) in section_headers:
            continue

        # Skip contact information.
        if "@" in clean:
            continue

        if re.search(
            r"\+?\d[\d\s().-]{6,}",
            clean
        ):
            continue

        # Skip URLs.
        if (
            "http://" in low
            or "https://" in low
            or "www." in low
        ):
            continue

        # Skip obvious location/contact lines.
        if any(
            keyword in low
            for keyword in [
                "jakarta, indonesia",
                "bandung, indonesia",
                "south jakarta",
                "indonesia |",
            ]
        ):
            continue

        # A candidate name is normally short.
        words = clean.split()

        if not (1 <= len(words) <= 4):
            continue

        # Must contain alphabetic characters.
        if not re.search(
            r"[A-Za-zÀ-ÿ]",
            clean
        ):
            continue

        # Avoid obvious section/title words.
        title_words = {
            "architect",
            "architectural",
            "drafter",
            "designer",
            "engineer",
            "developer",
            "intern",
            "freelancer",
            "profile",
            "summary",
            "about",
            "work",
            "experience",
            "education",
            "skills",
        }

        if all(
            word.lower() in title_words
            for word in words
        ):
            continue

        return clean

    return fallback

# ============================================================
# SKILL ONTOLOGY
#
# NOTE: This dictionary is treated as an optional "booster" for
# known synonym clusters. It is NOT the sole source of truth for
# what counts as a "skill" — that would only work for CVs in the
# architecture/design domain it was originally written for.
#
# Generic, domain-agnostic skill discovery happens separately
# via `extract_skill_tokens()`, which reads whatever the CV's
# own Skills section actually lists, regardless of industry.
# ============================================================

SKILL_ALIASES = {

    # Architecture software
    "autocad": [
        "autocad",
        "auto cad",
    ],

    "sketchup": [
        "sketchup",
        "sketch up",
    ],

    "revit": [
        "revit",
        "revit architecture",
    ],

    "lumion": [
        "lumion",
    ],

    "enscape": [
        "enscape",
    ],

    "v-ray": [
        "v-ray",
        "vray",
    ],

    "d5 render": [
        "d5 render",
        "d5",
    ],

    "twinmotion": [
        "twinmotion",
    ],

    "adobe photoshop": [
        "photoshop",
        "adobe photoshop",
    ],

    "adobe illustrator": [
        "illustrator",
        "adobe illustrator",
    ],

    # Architecture skills
    "architectural design": [
        "architectural design",
        "architecture design",
        "architectural design development",
    ],

    "interior design": [
        "interior design",
        "interior designer",
        "interior architecture",
        "interior architect",
    ],

    "technical drawing": [
        "technical drawing",
        "technical drawings",
        "architectural drawing",
        "architectural drawings",
        "working drawing",
        "working drawings",
        "shop drawing",
        "shop drawings",
        "as-built drawing",
        "as-built drawings",
        "blueprints",
    ],

    "3d modeling": [
        "3d modeling",
        "3d modelling",
        "3d model",
        "3d visualization",
        "3d visualisation",
        "3d rendering",
    ],

    "visualization": [
        "visualization",
        "visualisation",
        "rendering",
        "3d visualization",
        "3d visualisation",
    ],

    "construction management": [
        "construction management",
        "construction project",
        "construction projects",
    ],

    "project management": [
        "project management",
        "project coordination",
        "project coordination meetings",
    ],

    "site supervision": [
        "site supervision",
        "site supervisor",
        "supervised construction",
        "site inspection",
    ],

    "site survey": [
        "site survey",
        "site surveys",
        "site measurement",
        "site measurements",
        "site analysis",
    ],

    "building codes": [
        "building codes",
        "building code",
        "local building regulations",
        "building regulations",
        "regulatory standards",
        "regulations",
    ],

    "space planning": [
        "space planning",
        "space utilization",
        "space utilisation",
    ],

    "material specification": [
        "material specification",
        "material specifications",
        "material sourcing",
        "materials",
    ],

    # Generic professional skills
    "communication": [
        "communication",
        "communicating",
        "communicate effectively",
    ],

    "presentation": [
        "presentation",
        "client presentations",
        "presentations",
    ],

    "teamwork": [
        "teamwork",
        "team",
        "collaboration",
        "collaborate",
        "collaborating",
    ],

    "problem solving": [
        "problem solving",
        "problem-solving",
        "solve design challenges",
    ],

    "critical thinking": [
        "critical thinking",
    ],

    "attention to detail": [
        "attention to detail",
        "precision",
        "accuracy",
    ],

    "time management": [
        "time management",
        "project timeline",
        "project timelines",
    ],

    "creativity": [
        "creative",
        "creativity",
        "creative skills",
        "innovative designs",
    ],

    "visualization skills": [
        "visualization skills",
        "visualisation skills",
        "visualization",
    ],
}


def _contains_alias(text: str, alias: str) -> bool:
    """
    Safer matching for multi-word skills.
    """
    pattern = r"(?<!\w)" + re.escape(alias.lower()) + r"(?!\w)"
    return re.search(pattern, text.lower()) is not None


def _find_alias_span(text: str, alias: str):
    """
    Return the literal substring in `text` that matches `alias`,
    preserving the original casing/spacing found in the document
    (e.g. searching for the alias "teknik informatika" against a
    CV that literally wrote "Teknik Informatika" returns that
    exact original substring). Returns None if not found.
    """
    pattern = r"(?<!\w)" + re.escape(alias) + r"(?!\w)"
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(0) if match else None


def extract_skills(text: str):
    """
    Extract skills known via the curated SKILL_ALIASES booster.

    This intentionally stays narrow/curated. For a domain-agnostic
    view of "whatever the candidate says they can do", use
    `extract_skill_tokens()` on the CV's own Skills section instead.
    """
    found = []

    for canonical, aliases in SKILL_ALIASES.items():
        for alias in aliases:
            if _contains_alias(text, alias):
                found.append(canonical)
                break

    return sorted(set(found))


# ============================================================
# GENERIC / DOMAIN-AGNOSTIC SKILL TOKEN EXTRACTION
# ============================================================

_SKILL_TOKEN_STOPWORDS = {
    "dan", "and", "atau", "or", "dengan", "with", "serta",
    "including", "termasuk", "etc", "dll", "lainnya", "other",
    "others", "baik", "seperti", "such", "as", "for", "untuk",
    "yang", "di", "dalam", "in", "of", "the", "a", "an",
}


def extract_skill_tokens(section_text: str):
    """
    Extract free-form skill tokens directly from a CV's Skills
    section (or similar), independent of any predefined skill
    dictionary. This is what makes matching work for CVs from
    ANY industry, not just architecture/design.

    A CV that lists:
        "Recruitment, HRIS, Payroll Processing, Labor Law,
         Employee Relations, Ms. Office"

    will yield those exact tokens (lowercased, cleaned) instead
    of returning nothing just because they aren't in a hardcoded
    architecture-focused dictionary.
    """

    if not section_text or not section_text.strip():
        return []

    tokens = []

    for line in section_text.splitlines():

        line = line.strip()
        if not line:
            continue

        # Strip common bullet / list markers.
        line = re.sub(r"^[\-•●▪◦*·»›\u2022]+\s*", "", line)

        # A skills section is usually delimited by commas,
        # bullets, pipes, semicolons, or slashes rather than
        # being full prose sentences.
        parts = re.split(r"[,;|/•●▪◦]", line)

        for part in parts:

            part = part.strip(" .\t")
            part = re.sub(r"\s+", " ", part)

            if not part:
                continue

            words = part.split()

            # Keep only short, list-like phrases (real skill
            # labels), not full sentences that slipped in.
            if not (1 <= len(words) <= 5):
                continue

            low = part.lower()

            if low in _SKILL_TOKEN_STOPWORDS:
                continue

            if not re.search(r"[A-Za-zÀ-ÿ]", part):
                continue

            tokens.append(low)

    return sorted(set(tokens))


# ============================================================
# EDUCATION EXTRACTION
# ============================================================

DEGREE_ALIASES = {
    "bachelor": [
        "bachelor",
        "bachelors",
        "bachelor's",
        "sarjana",
        "s1",
        "b.arch",
    ],
    "master": [
        "master",
        "master's",
        "magister",
        "s2",
    ],
    "diploma": [
        "diploma",
        "d3",
        "d4",
    ],
    "vocational": [
        "vocational",
        "smk",
        "smkn",
    ],
}

# ============================================================
# EDUCATION FIELDS — booster dictionary
#
# Same idea as SKILL_ALIASES: kept as a fallback so common
# fields still resolve to a stable canonical bucket, but it is
# no longer the only way education fields are detected — and
# critically, the *canonical bucket name* (e.g. "information
# technology") is no longer what gets shown to the user. See
# `extract_education()` below: the LITERAL alias text that was
# actually found in the document (e.g. "teknik informatika") is
# what ends up in `fields_raw`/`fields_display`.
# ============================================================

EDUCATION_FIELDS = {
    "architecture": [
        "architecture",
        "architectural engineering",
        "arsitektur",
    ],

    "interior design": [
        "interior design",
        "interior designer",
        "interior architecture",
        "interior architect",
    ],

    "civil engineering": [
        "civil engineering",
        "teknik sipil",
        "sipil",
    ],

    "design": [
        "design",
        "desain",
    ],

    "information technology": [
        "information technology",
        "informatika",
        "computer science",
        "ilmu komputer",
        "teknik informatika",
        "sistem informasi",
        "information systems",
    ],

    "human resources": [
        "human resources",
        "human resource management",
        "manajemen sumber daya manusia",
        "msdm",
        "psikologi",
        "psychology",
    ],

    "business management": [
        "business administration",
        "management",
        "manajemen",
        "administrasi bisnis",
        "marketing",
        "pemasaran",
    ],

    "accounting finance": [
        "accounting",
        "akuntansi",
        "finance",
        "keuangan",
        "ekonomi",
        "economics",
    ],

    "communication": [
        "communication",
        "communication science",
        "ilmu komunikasi",
        "public relations",
    ],

    "law": [
        "law",
        "hukum",
        "ilmu hukum",
    ],
}


# Regex patterns that capture the *literal* field name mentioned
# in text, e.g. "Bachelor of Computer Science", "S1 Teknik
# Informatika", "Jurusan Manajemen", "Degree in Marketing".
_EDUCATION_FIELD_PATTERNS = [
    # English style: "Bachelor of/in X", "Master's in X",
    # "Sarjana di X", "S1 jurusan X"
    r"(?:bachelor'?s?|master'?s?|diploma|sarjana|magister)"
    r"\s+(?:degree\s+)?(?:in|of|dalam|di|jurusan)\s+"
    r"([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s/&\-]{2,60})",

    # Indonesian short degree codes directly followed by the
    # field name, with or without a connector word/dash —
    # e.g. "S1 Teknik Informatika", "S1 - Manajemen",
    # "D3 Akuntansi". This is the common format that was
    # previously MISSED, causing the booster dictionary
    # (canonical bucket name) to be shown instead of the
    # literal field the CV actually wrote.
    r"\b(?:s1|s2|d3|d4)\b\s*[-:]?\s*"
    r"(?:jurusan\s+|di\s+|dalam\s+)?"
    r"([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s/&\-]{2,60})",

    r"(?:jurusan|program studi|prodi)\s+"
    r"([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s/&\-]{2,60})",

    r"degree\s+in\s+"
    r"([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s/&\-]{2,60})",
]

_EDUCATION_FIELD_STOPWORDS = {
    "a", "an", "the", "related", "field", "fields", "etc",
    "other", "others", "any", "relevant", "similar",
}

# If a captured "field" phrase actually starts with one of these
# words, it's almost certainly an institution name that got
# swept up by the adjacency pattern above (e.g. "S1 Universitas
# XYZ"), not a field of study — so it gets discarded rather than
# shown as if it were the candidate's major.
_INSTITUTION_LEAD_WORDS = {
    "universitas", "university", "institut", "institute",
    "politeknik", "polytechnic", "sekolah", "college", "akademi",
    "academy", "stikes", "stie", "stmik", "stia", "upi", "itb",
    "ugm", "ui", "its",
}


def _looks_like_institution(phrase: str) -> bool:
    words = phrase.lower().split()
    return bool(words) and words[0] in _INSTITUTION_LEAD_WORDS


def extract_education_field_phrases(text: str):
    """
    Pull the literal field-of-study phrase(s) mentioned in a
    piece of text (a JD requirement line OR a CV education
    entry), e.g.:

        "Degree in architecture, interior design, or a
         related field"
        -> ["architecture", "interior design"]

        "S1 Teknik Informatika"
        -> ["teknik informatika"]

    This works for ANY field name, not just the ones hardcoded
    in EDUCATION_FIELDS.
    """

    if not text:
        return []

    phrases = []

    for pattern in _EDUCATION_FIELD_PATTERNS:

        for match in re.finditer(pattern, text, flags=re.IGNORECASE):

            chunk = match.group(1)

            # Drop trailing "...or a related field" style noise.
            chunk = re.sub(
                r"\bor\s+(?:a\s+)?related\s+field(s)?\b.*",
                "",
                chunk,
                flags=re.IGNORECASE,
            )
            chunk = re.sub(
                r"\brelated\s+field(s)?\b",
                "",
                chunk,
                flags=re.IGNORECASE,
            )

            parts = re.split(r",|/|\band\b|\bor\b|&", chunk, flags=re.IGNORECASE)

            for part in parts:

                part = part.strip(" .\t")
                part = re.sub(r"\s+", " ", part)

                if not part:
                    continue

                words = part.split()

                if not (1 <= len(words) <= 5):
                    continue

                if part.lower() in _EDUCATION_FIELD_STOPWORDS:
                    continue

                if _looks_like_institution(part):
                    continue

                phrases.append(part.lower())

    return list(dict.fromkeys(phrases))


def extract_education(text: str):
    """
    Extract highest / most relevant education level and
    education fields from the education section.

    Returns:
        degree:         "bachelor" | "master" | "diploma" |
                         "vocational" | None
        fields:          canonical bucket names (e.g.
                         "information technology") — used
                         internally for reliable matching.
        fields_raw:      literal phrase(s) as actually written
                         in the document (e.g. "teknik
                         informatika") — booster-dictionary
                         alias hits AND dynamically regex-
                         extracted phrases, combined.
        fields_display:  best human-readable version for UI —
                         prefers fields_raw (what the candidate
                         actually wrote), falls back to the
                         canonical bucket names only if nothing
                         literal was captured.
    """

    low = normalize_text(text)

    # --------------------------------------------------------
    # DEGREE
    # --------------------------------------------------------

    degree_priority = [
        ("master", [
            "master",
            "master's",
            "magister",
            "s2",
        ]),

        ("bachelor", [
            "bachelor",
            "bachelors",
            "bachelor's",
            "sarjana",
            "s1",
            "b.arch",
            "bachelor of architecture",
        ]),

        ("diploma", [
            "diploma",
            "d3",
            "d4",
        ]),

        ("vocational", [
            "vocational",
            "smk",
            "smkn",
        ]),
    ]

    degree = None

    for canonical, aliases in degree_priority:

        if any(
            _contains_alias(low, alias)
            for alias in aliases
        ):
            degree = canonical
            break

    # --------------------------------------------------------
    # UNIVERSITY DETECTION
    # --------------------------------------------------------

    university_indicators = [
        "university",
        "universitas",
        "institute",
        "institut",
        "college",
    ]

    has_university = any(
        _contains_alias(low, indicator)
        for indicator in university_indicators
    )

    # If the candidate has university education but
    # the CV does not explicitly write "Bachelor",
    # treat it as bachelor-level education for this POC.
    if degree == "vocational" and has_university:
        degree = "bachelor"

    # --------------------------------------------------------
    # EDUCATION FIELD (booster dictionary — canonical buckets
    # + the literal alias text that actually matched, so the
    # UI can show what the CV really says instead of a
    # translated/renamed bucket label).
    # --------------------------------------------------------

    fields = []
    fields_from_booster_literal = []

    for canonical, aliases in EDUCATION_FIELDS.items():

        matched_aliases = [
            alias
            for alias in aliases
            if _contains_alias(low, alias)
        ]

        if not matched_aliases:
            continue

        fields.append(canonical)

        # Prefer the longest/most specific alias that matched
        # (e.g. "teknik informatika" over just "informatika")
        # so the displayed phrase is as complete as possible.
        best_alias = max(matched_aliases, key=len)
        literal = _find_alias_span(text, best_alias)

        if literal:
            fields_from_booster_literal.append(literal.strip().lower())

    # --------------------------------------------------------
    # EDUCATION FIELD (dynamic — literal phrase from the text,
    # for fields that aren't in the booster dictionary at all,
    # e.g. an unusual major name).
    # --------------------------------------------------------

    fields_raw_dynamic = extract_education_field_phrases(text)

    fields_raw = list(dict.fromkeys(
        fields_from_booster_literal + fields_raw_dynamic
    ))

    fields_display = fields_raw if fields_raw else sorted(set(fields))

    return {
        "degree": degree,
        "fields": sorted(set(fields)),
        "fields_raw": fields_raw,
        "fields_display": fields_display,
    }


# ============================================================
# EXPERIENCE EXTRACTION
# ============================================================

MONTHS = {
    # English
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,

    # Indonesian
    "januari": 1,
    "februari": 2,
    "maret": 3,
    "mei": 5,
    "juni": 6,
    "juli": 7,
    "agustus": 8,
    "oktober": 10,
    "desember": 12,
}

PRESENT_WORDS = {
    "present",
    "currently",
    "current",
    "now",
    "sekarang",
    "saat ini",
    "hingga sekarang",
}


def _parse_month_year(month, year):
    month = month.lower().strip()

    if month not in MONTHS:
        return None

    return date(
        int(year),
        MONTHS[month],
        1
    )


def _parse_date_token(token):
    """
    Parse common CV date tokens.

    Supported:
        2024
        Jan 2024
        January 2024
        2024 Jan
        06/2024
        06-2024
        2024/06
        2024-06
    """

    token = token.strip().lower()
    token = re.sub(r"\s+", " ", token)

    # YYYY
    if re.fullmatch(r"\d{4}", token):
        return date(int(token), 1, 1)

    # MM/YYYY or MM-YYYY
    match = re.fullmatch(r"(\d{1,2})[/-](\d{4})", token)
    if match:
        month = int(match.group(1))
        year = int(match.group(2))
        if 1 <= month <= 12:
            return date(year, month, 1)

    # YYYY/MM or YYYY-MM
    match = re.fullmatch(r"(\d{4})[/-](\d{1,2})", token)
    if match:
        year = int(match.group(1))
        month = int(match.group(2))
        if 1 <= month <= 12:
            return date(year, month, 1)

    # Month YYYY
    match = re.fullmatch(r"([a-zA-Z]+)\s+(\d{4})", token)
    if match:
        return _parse_month_year(match.group(1), match.group(2))

    # YYYY Month
    match = re.fullmatch(r"(\d{4})\s+([a-zA-Z]+)", token)
    if match:
        return _parse_month_year(match.group(2), match.group(1))

    return None


def _parse_end_date(token):
    token = token.strip().lower()

    if token in PRESENT_WORDS:
        return date.today().replace(day=1)

    return _parse_date_token(token)


def _months_between(start, end):
    return (
        (end.year - start.year) * 12
        + (end.month - start.month)
        + 1
    )


def _extract_date_ranges(text):
    """
    Extract employment date ranges from CV text.

    Supported examples:
        2023 - Present
        2021 - 2024
        Jan 2022 - Present
        January 2022 - December 2024
        Jun 2018 - Jul 2019
        2018 - 2023
        06/2022 - Present
        06/2022 - 12/2024
        2022 - Sekarang
        2022 - Saat ini
        2022 - Hingga sekarang

    Unsupported/malformed patterns are ignored instead of
    crashing the CV parser.
    """

    if not text:
        return []

    text = str(text)
    ranges = []
    month_names = list(MONTHS.keys())

    # 1. Month Year - Month Year
    pattern_month_month = (
        r"\b([A-Za-z]+)\s+(\d{4})"
        r"\s*[-–—]\s*"
        r"([A-Za-z]+)\s+(\d{4})\b"
    )

    for match in re.finditer(pattern_month_month, text, flags=re.IGNORECASE):
        start = _parse_date_token(f"{match.group(1)} {match.group(2)}")
        end = _parse_date_token(f"{match.group(3)} {match.group(4)}")

        if start is not None and end is not None and end >= start:
            ranges.append((start, end))

    # 2. Month Year - Present
    pattern_month_present = (
        r"\b([A-Za-z]+)\s+(\d{4})"
        r"\s*[-–—]\s*"
        r"(Present|Currently|Current|Now|Sekarang|Saat ini|Hingga sekarang)\b"
    )

    for match in re.finditer(pattern_month_present, text, flags=re.IGNORECASE):
        start = _parse_date_token(f"{match.group(1)} {match.group(2)}")
        end = _parse_end_date(match.group(3))

        if start is not None and end is not None and end >= start:
            ranges.append((start, end))

    # 3. Numeric Month/Year - Numeric Month/Year
    pattern_numeric_month_month = (
        r"\b(\d{1,2}[/-]\d{4})"
        r"\s*[-–—]\s*"
        r"(\d{1,2}[/-]\d{4})\b"
    )

    for match in re.finditer(
        pattern_numeric_month_month, text, flags=re.IGNORECASE
    ):
        start = _parse_date_token(match.group(1))
        end = _parse_date_token(match.group(2))

        if start is not None and end is not None and end >= start:
            ranges.append((start, end))

    # 4. Numeric Month/Year - Present
    pattern_numeric_month_present = (
        r"\b(\d{1,2}[/-]\d{4})"
        r"\s*[-–—]\s*"
        r"(Present|Currently|Current|Now|Sekarang|Saat ini|Hingga sekarang)\b"
    )

    for match in re.finditer(
        pattern_numeric_month_present, text, flags=re.IGNORECASE
    ):
        start = _parse_date_token(match.group(1))
        end = _parse_end_date(match.group(2))

        if start is not None and end is not None and end >= start:
            ranges.append((start, end))

    # 5. Year - Present
    pattern_year_present = (
        r"(?<![A-Za-z])"
        r"(\d{4})"
        r"\s*[-–—]\s*"
        r"(Present|Currently|Current|Now|Sekarang|Saat ini|Hingga sekarang)\b"
    )

    for match in re.finditer(
        pattern_year_present, text, flags=re.IGNORECASE
    ):
        before = text[max(0, match.start() - 30):match.start()].lower()

        # Hindari mengambil bagian "2021 - Present"
        # dari "November 2021 - Present".
        if any(
            re.search(rf"\b{re.escape(month)}\s+$", before)
            for month in month_names
        ):
            continue

        start = _parse_date_token(match.group(1))
        end = _parse_end_date(match.group(2))

        if start is not None and end is not None and end >= start:
            ranges.append((start, end))

    # 6. Year - Year
    pattern_year_year = (
        r"\b(\d{4})"
        r"\s*[-–—]\s*"
        r"(\d{4})\b"
    )

    for match in re.finditer(pattern_year_year, text, flags=re.IGNORECASE):
        before = text[max(0, match.start() - 30):match.start()].lower()

        # Hindari double-counting "2021 - 2024"
        # dari "January 2021 - December 2024".
        if any(
            re.search(rf"\b{re.escape(month)}\s+$", before)
            for month in month_names
        ):
            continue

        start = _parse_date_token(match.group(1))
        end = _parse_date_token(match.group(2))

        if start is not None and end is not None and end >= start:
            ranges.append((start, end))

    # 7. Remove duplicates
    unique_ranges = []
    for item in ranges:
        if item not in unique_ranges:
            unique_ranges.append(item)

    return unique_ranges

def _merge_date_ranges(ranges):
    """
    Merge overlapping employment periods so that
    overlapping jobs are not double-counted.
    """

    if not ranges:
        return []

    ranges = sorted(
        ranges,
        key=lambda item: item[0]
    )

    merged = [
        list(ranges[0])
    ]

    for start, end in ranges[1:]:

        previous_start, previous_end = merged[-1]

        if start <= previous_end:

            if end > previous_end:
                merged[-1][1] = end

        else:
            merged.append(
                [start, end]
            )

    return [
        (start, end)
        for start, end in merged
    ]


# Words that, when found shortly before/after a "<number> tahun"
# phrase, mean the number is almost certainly NOT years of work
# experience (age, contract length, certificate validity, etc.).
# This fixes a bug where e.g. "Usia: 24 tahun" was being read as
# 24 years of professional experience.
_TAHUN_EXCLUDE_CONTEXT = [
    "usia", "umur", "age",
    "kontrak", "contract",
    "berlaku", "valid", "validity",
    "garansi", "warranty",
    "sertifikat", "certificate", "license", "lisensi",
]


def extract_experience_years(text):
    """
    Extract total professional experience.

    Priority:
    1. Explicit statements such as:
       "5 years of experience"
    2. Employment date ranges:
       "2021 - Present"
       "Jan 2022 - Dec 2024"

    The whole CV is used instead of relying only on the
    EXPERIENCE section because CV layouts vary heavily.
    """

    if not text:
        return 0.0

    # --------------------------------------------------------
    # 1. Explicit experience statements
    # --------------------------------------------------------

    explicit_values = []

    explicit_patterns = [
        r"(\d+(?:\.\d+)?)\s*\+?\s*"
        r"(?:years?|yrs?)\s+(?:of\s+)?experience",

        r"(?:experience|pengalaman)"
        r"\s*[:\-]?\s*"
        r"(\d+(?:\.\d+)?)\s*"
        r"(?:years?|yrs?|tahun)",

        r"with\s+(\d+(?:\.\d+)?)\s*"
        r"(?:years?|yrs?)",
    ]

    for pattern in explicit_patterns:

        matches = re.findall(
            pattern,
            text,
            flags=re.IGNORECASE
        )

        for value in matches:

            try:
                explicit_values.append(
                    float(value)
                )
            except ValueError:
                pass

    # --------------------------------------------------------
    # 1b. "<number> tahun (pengalaman)" — handled separately
    # with context-aware exclusion so age / contract length /
    # certificate validity etc. don't get misread as work
    # experience.
    # --------------------------------------------------------

    tahun_pattern = r"(\d+(?:\.\d+)?)\s*tahun\b"

    for match in re.finditer(tahun_pattern, text, flags=re.IGNORECASE):

        window_before = text[max(0, match.start() - 25):match.start()].lower()
        window_after = text[match.end():match.end() + 25].lower()

        if any(word in window_before for word in _TAHUN_EXCLUDE_CONTEXT):
            continue
        if any(word in window_after for word in _TAHUN_EXCLUDE_CONTEXT):
            continue

        # Only trust a bare "<number> tahun" as work experience
        # if "pengalaman"/"kerja" is nearby — otherwise it's too
        # ambiguous (could be age, duration of a course, etc.).
        has_experience_context = (
            "pengalaman" in window_before
            or "pengalaman" in window_after
            or "kerja" in window_before
            or "kerja" in window_after
        )

        if not has_experience_context:
            continue

        try:
            explicit_values.append(float(match.group(1)))
        except ValueError:
            pass

    # --------------------------------------------------------
    # 2. Date-based experience
    # --------------------------------------------------------

    ranges = _extract_date_ranges(
        text
    )

    if ranges:

        merged = _merge_date_ranges(
            ranges
        )

        total_months = sum(
            _months_between(
                start,
                end
            )
            for start, end in merged
        )

        date_based_years = (
            total_months / 12.0
        )

    else:
        date_based_years = 0.0

    # --------------------------------------------------------
    # 3. Combine evidence
    # --------------------------------------------------------

    candidates = [
        date_based_years,
        *explicit_values
    ]

    if not candidates:
        return 0.0

    result = max(candidates)

    return round(
        result,
        1
    )

# ============================================================
# EXPERIENCE TEXT
# ============================================================

def extract_experience_section(text: str):
    """
    Extract professional/work experience section.

    Supports both English and Indonesian CV section names.
    """

    return _section(
        text,
        [
            "experience",
            "experiences",
            "work experience",
            "work experiences",
            "professional experience",
            "professional experiences",
            "employment history",
            "career history",
            "work history",

            "pengalaman",
            "pengalaman kerja",
            "riwayat kerja",
            "riwayat karier",
        ]
    )


# ============================================================
# FULL CV PARSER
# ============================================================

def parse_cv(text: str, source=""):
    """
    Convert raw CV text into a structured candidate profile.
    """

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    experience_text = extract_experience_section(
        text
    )

    # Hanya gunakan experience section untuk menghitung
    # pengalaman kerja.
    experience_years = extract_experience_years(
        experience_text
    )

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    skills_text = _section(
        text,
        [
            "skills",
            "skill",
            "technical skills",
            "technical skill",
            "soft skills",
            "soft skill",
            "hard skills",
            "hard skill",
            "professional skills",
            "core skills",

            "keahlian",
            "keahlian teknis",
            "kemampuan",
            "keterampilan",
        ]
    )

    # Curated booster (known synonym clusters).
    skills = extract_skills(
        text
    )

    # Domain-agnostic: whatever the CV's own Skills section
    # actually lists, regardless of industry.
    skills_raw = extract_skill_tokens(
        skills_text
    )

    # Combined view used for matching — union of both sources.
    all_skills = sorted(set(skills) | set(skills_raw))

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    # Prioritize FORMAL EDUCATION because some CVs use
    # multi-column layouts. PDF text extraction may place
    # another column between EDUCATION and its actual content.
    education_text = _section(
        text,
        [
            "formal education",
        ]
    )

    # Fallback for CVs that do not have a FORMAL EDUCATION
    # subsection.
    if not education_text:
        education_text = _section(
            text,
            [
                "education",
                "education level",
                "academic background",
                "academic qualification",
                "academic qualifications",
                "educational background",
                "educational qualification",
                "education qualification",

                "pendidikan",
                "riwayat pendidikan",
                "latar belakang pendidikan",
            ]
        )

    education = extract_education(
        education_text
)

    # --------------------------------------------------------
    # RETURN
    # --------------------------------------------------------

    return {
        "name": extract_name(
            text,
            Path(source).stem
            if source
            else "Unknown Candidate"
        ),

        "skills": all_skills,

        "skills_curated": skills,

        "skills_raw": skills_raw,

        "experience_years":
            experience_years,

        "education":
            education,

        "experience_text":
            experience_text,

        "skills_text":
            skills_text,

        "education_text":
            education_text,

        "raw_text":
            text,

        "source":
            source,
    }